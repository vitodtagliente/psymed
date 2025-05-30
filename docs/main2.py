import spacy
import sys
from docx import Document

# Tentativo di caricamento del modello italiano di spaCy
try:
    nlp = spacy.load("it_core_news_sm")
except OSError:
    sys.exit(
        "Modello 'it_core_news_sm' non trovato.\n" +
        "Installa il modello con: python -m spacy download it_core_news_sm"
    )

# Lessico BPRS: mappa item -> lista di lemmi/termini chiave (in italiano)
LEXICON = {
    "preoccupazione somatica": ["somatico", "dolore", "malessere", "distress", "ipocondria", "paura malattia", "salute fisica"],
    "ansia": ["ansia", "tensione", "paura", "agitazione", "preoccupazione", "nervosismo", "inquietudine", "apprensione"],
    "ritiro emotivo": ["isolamento", "ritiro", "distacco", "mancanza interazione", "relazione", "difficoltà a relazionarsi", "isolamento sociale"],
    "disorganizzazione concettuale": ["pensiero confuso", "disconnesso", "disorganizzato", "pensieri frammentati", "illogico", "difficoltà a seguire un discorso"],
    "sentimenti di colpa": ["senso di colpa", "vergogna", "rimorso", "auto-rimprovero", "auto-accusa"],
    "tensione": ["tensione fisica", "tensione motoria", "nervosismo", "iperattivazione", "irrequietezza", "stato di allerta"],
    "manierismi e posture": ["manierismi", "posture insolite", "bizzarri", "innaturali", "comportamento motorio peculiare", "stereotipie"],
    "grandiosità": ["esagerata opinione di sé", "arroganza", "convinzione di poteri insoliti", "abilità superiori", "delusioni di grandezza"],
    "umore depressivo": ["tristezza", "malinconia", "abbattimento", "pessimismo", "disperazione", "depressione", "umore nero"],
    "ostilità": ["animosità", "disprezzo", "belligeranza", "disprezzo per gli altri", "aggressività", "rabbia", "risentimento"],
    "sospettosità": ["diffidenza", "sospetto", "paranoia", "credenza di intenzioni maliziose", "persecuzione", "sentirsi spiati"],
    "comportamento allucinatorio": ["allucinazioni", "percezioni senza stimolo esterno", "voci", "visioni", "sensazioni insolite"],
    "rallentamento motorio": ["movimenti rallentati", "parola rallentata", "ridotto tono corporeo", "apatia motoria", "bradicinesia"],
    "non collaborazione": ["resistenza", "chiusura", "rifiuto dell'autorità", "mancanza di cooperazione", "ostinazione"],
    "contenuto ideativo inusuale": ["pensieri strani", "insoliti", "bizzarri", "delusioni", "idee fisse", "pensiero magico"],
    "affettività appiattita": ["tono emotivo ridotto", "mancanza di sentimenti", "freddezza emotiva", "appiattimento emotivo", "espressione facciale ridotta"],
    "eccitazione": ["tono emotivo intensificato", "agitazione", "aumentata reattività", "iperattività", "euforia", "irritabilità"],
    "disorientamento": ["disorientamento spazio-temporale", "confusione", "mancanza di consapevolezza della persona, del luogo o del tempo"]
}

# Range punteggi per ogni item (min, max)
ITEMS_RANGE = {item: (1, 7) for item in LEXICON}

def score_bprs(doc):
    """Calcola i punteggi per ciascun item BPRS e il totale."""
    scores = {item: 1 for item in LEXICON}
    for token in doc:
        for item, terms in LEXICON.items():
            if token.lemma_ in terms:
                min_s, max_s = ITEMS_RANGE[item]
                if scores[item] < max_s:
                    scores[item] += 1
    total = sum(scores.values())
    return scores, total

def estimate_gaf(bprs_total: int, n_items: int):
    """Stima il GAF in scala 0-100, inversamente proporzionale al BPRS."""
    min_total = n_items * 1
    max_total = n_items * 7
    gaf = 100 - ((bprs_total - min_total) / (max_total - min_total)) * 100
    return round(gaf, 1)

def read_docx_to_text(file_path):
    """
    Reads a .docx file and returns its full text content as a string.
    Includes text from paragraphs and tables.
    """
    try:
        doc = Document(file_path)
        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        return '\n'.join(text_parts)

    except Exception as e:
        return f"Error reading DOCX file: {e}"

def read_to_text(file_path):
    try:
        with open(file_path, encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        sys.exit(f"File non trovato: {file_path}")
        return ""

def main():
    # text = read_to_text("referto.txt")
    text = read_docx_to_text('ubaldini_ubaldo.docx')

    doc = nlp(text)
    print("--- Risultati dell'analisi delle entità ---")
    for ent in doc.ents:
        print(f"Testo: {ent.text}, Categoria: {ent.label_}")
    print("")

    scores, total = score_bprs(doc)
    gaf = estimate_gaf(total, len(scores))

    print("=== BPRS ===")
    for item, score in scores.items():
        print(f"{item}: {score}")
    print(f"Punteggio totale BPRS: {total}\n")
    print("=== Stima GAF ===")
    print(f"GAF stimato: {gaf}")

    print("")

if __name__ == "__main__":
    main()
